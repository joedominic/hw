"""
Django Ninja API: resume optimization, LLM settings, workflows, prompts, uploads.

Mounted with the app’s `/api/resume/` prefix. Job search and pipeline JSON live in `jobs_api`.
"""
from ninja import Router, File, Schema, Form
from ninja.files import UploadedFile
from ninja.errors import HttpError
from django.shortcuts import get_object_or_404
from django.http import FileResponse
from django.utils import timezone
from .models import (
    AtsJudgeProfile,
    UserResume,
    JobDescription,
    OptimizedResume,
    AgentLog,
    LLMProviderConfig,
    OptimizerWorkflow,
)
from .jobs_api import router as jobs_router
from .tasks import optimize_resume_task
from .prompts import DEFAULT_MATCHING_PROMPT
from langchain_core.messages import HumanMessage, SystemMessage

from .agents import (
    DEFAULT_WRITER_PROMPT,
    DEFAULT_ATS_JUDGE_PROMPT,
    DEFAULT_RECRUITER_JUDGE_PROMPT,
    DEFAULT_FIT_CHECK_PROMPT,
    VALID_STEP_IDS,
    get_llm,
    run_fit_check,
    writer_node,
    ats_judge_node,
    recruiter_judge_node,
)
from .llm_gateway import (
    call_invoke_llm_messages,
    LLMRequestsDisabled,
    USAGE_QUERY_API_LLM_COMPLETE,
    USAGE_QUERY_API_RESUME_FIT,
)
from .services import parse_pdf
from .crypto import encrypt_api_key, decrypt_api_key
from .llm_services import list_models_for_provider, is_auth_error, DEFAULT_MODELS, LLM_PROVIDERS
from typing import List, Optional
from django.conf import settings
import json
import io
import tempfile
import os
import logging
import traceback

logger = logging.getLogger(__name__)

router = Router()
router.add_router("/jobs", jobs_router)

MAX_JOB_DESCRIPTION_LENGTH = 50_000
MAX_RESUME_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
ALLOWED_CONTENT_TYPES = ("application/pdf",)
PDF_EXTENSION = ".pdf"


def _require_api_auth(request):
    """
    Simple optional auth: if settings.API_ACCESS_TOKEN is set, require
    header X-Api-Token to match. If not set, no auth is enforced.
    """
    token = getattr(settings, "API_ACCESS_TOKEN", None)
    if not token:
        return
    provided = request.headers.get("X-Api-Token") or request.META.get("HTTP_X_API_TOKEN")
    if not provided or provided != token:
        raise HttpError(401, "Missing or invalid API access token.")

class OptimizeRequest(Schema):
    job_description: str
    llm_provider: str
    llm_model: Optional[str] = None  # from provider's list; default from config or built-in
    api_key: Optional[str] = None  # optional when key stored in DB or env
    prompt_writer: Optional[str] = None  # override default writer prompt template
    prompt_ats_judge: Optional[str] = None  # override default ATS judge prompt template
    prompt_recruiter_judge: Optional[str] = None  # override default recruiter judge prompt template
    ats_judge_profile_id: Optional[int] = None  # named ATS profile from library
    optimizer_workflow_id: Optional[int] = None  # saved workflow (steps + default ATS)
    debug: Optional[bool] = None  # when True, log full prompts sent to LLM in agent logs
    workflow_steps: Optional[str] = None  # JSON array of step ids, e.g. ["writer","ats_judge","recruiter_judge"]
    loop_to: Optional[str] = None  # step to loop back to after last step; empty = single pass
    score_threshold: Optional[int] = None  # exit when avg score >= this (0-100, default 85)
    optimization_notes: Optional[str] = None
    pipeline_skills_json: Optional[str] = None
    job_highlights: Optional[str] = None

class StatusResponse(Schema):
    status: str
    status_display: Optional[str] = None
    ats_score: Optional[int] = None
    recruiter_score: Optional[int] = None
    optimized_content: Optional[str] = None
    error_message: Optional[str] = None
    total_input_tokens: Optional[int] = None
    total_output_tokens: Optional[int] = None
    logs: List[dict] = []
    optimizer_context_snapshot: Optional[dict] = None

class PromptsResponse(Schema):
    writer: str
    writer_system: str = ""
    writer_user: str = ""
    ats_judge: str
    ats_judge_system: str = ""
    ats_judge_user: str = ""
    recruiter_judge: str
    recruiter_judge_system: str = ""
    recruiter_judge_user: str = ""
    matching: str
    matching_system: str = ""
    matching_user: str = ""
    insights: str
    insights_system: str = ""
    insights_user: str = ""
    jd_cleanse: str
    jd_cleanse_system: str = ""
    jd_cleanse_user: str = ""

class FitCheckRequest(Schema):
    job_description: str
    llm_provider: str
    llm_model: Optional[str] = None
    api_key: Optional[str] = None
    prompt_fit_check: Optional[str] = None

class FitCheckResponse(Schema):
    score: int
    reasoning: str
    thoughts: str = ""


class RunStepRequest(Schema):
    step: str  # writer | ats_judge | recruiter_judge
    job_description: str
    use_resume_id: Optional[int] = None
    optimized_resume: Optional[str] = None  # draft text; writer uses when non-empty. Judges: draft or file/use_resume_id (PDF text).
    feedback: Optional[str] = None  # for writer (comma-separated)
    prompt_writer: Optional[str] = None
    prompt_ats_judge: Optional[str] = None
    prompt_recruiter_judge: Optional[str] = None
    ats_judge_profile_id: Optional[int] = None
    optimizer_workflow_id: Optional[int] = None
    llm_provider: str
    llm_model: Optional[str] = None
    debug: Optional[bool] = None
    job_cache_key: Optional[str] = None
    optimization_notes: Optional[str] = None
    pipeline_skills_json: Optional[str] = None
    job_highlights: Optional[str] = None


class RunStepResponse(Schema):
    """Response for POST /run-step. Frontend: use output.debug_prompt for Input, output for Output, output.input_tokens/output_tokens for token count. error is set when the step failed."""
    step: str  # "writer" | "ats_judge" | "recruiter_judge"
    output: dict  # writer: {optimized_resume, ...}; ats/recruiter: {scores, feedback, ...}; judges accept draft or PDF via same file/use_resume_id as Writer when draft empty
    error: Optional[str] = None  # when set, step failed; frontend should show step name + error


class LlmCompleteRequest(Schema):
    user: str
    system: Optional[str] = None
    job_cache_key: Optional[str] = None
    llm_provider: Optional[str] = None
    llm_model: Optional[str] = None


class LlmCompleteResponse(Schema):
    content: str
    input_tokens: Optional[int] = None
    output_tokens: Optional[int] = None


@router.post("/llm/complete", response=LlmCompleteResponse)
def llm_complete(request, payload: LlmCompleteRequest):
    """Invoke the central LLM gateway with system + user strings (optional provider override)."""
    _require_api_auth(request)
    u_text = (payload.user or "").strip()
    if not u_text:
        raise HttpError(400, "user is required")
    messages = []
    if (payload.system or "").strip():
        messages.append(SystemMessage(content=payload.system.strip()))
    messages.append(HumanMessage(content=u_text))
    llm_override = None
    if payload.llm_provider:
        prov = payload.llm_provider.strip()
        if prov not in LLM_PROVIDERS:
            raise HttpError(400, f"llm_provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")
        config = LLMProviderConfig.objects.filter(provider=prov).first()
        if not config or not config.encrypted_api_key:
            raise HttpError(400, f"API key required for {prov}.")
        api_key = decrypt_api_key(config.encrypted_api_key)
        model = payload.llm_model or config.default_model or None
        llm_override = get_llm(prov, api_key, model)
    try:
        raw = call_invoke_llm_messages(
            messages,
            job_cache_key=(payload.job_cache_key or "").strip() or None,
            llm_override=llm_override,
            usage_query_kind=USAGE_QUERY_API_LLM_COMPLETE,
        )
    except LLMRequestsDisabled as e:
        raise HttpError(503, str(e)) from e
    from .agents import _normalize_token_usage

    content = getattr(raw, "content", None) if raw is not None else None
    if content is None and raw is not None:
        content = str(raw)
    content = content or ""
    u = _normalize_token_usage(raw, getattr(raw, "llm_output", None), None) if raw is not None else {}
    return LlmCompleteResponse(
        content=content if isinstance(content, str) else str(content),
        input_tokens=u.get("input_tokens"),
        output_tokens=u.get("output_tokens"),
    )


@router.post("/run-step", response=RunStepResponse)
def run_step(
    request,
    payload: RunStepRequest = Form(...),
    file: Optional[UploadedFile] = File(None),
):
    """Run a single optimizer step (writer, ats_judge, or recruiter_judge) and return its output."""
    step = (payload.step or "").strip().lower()
    logger.warning("[run_step] step=%s, job_desc_len=%s, optimized_resume_len=%s", step, len(payload.job_description or ""), len(payload.optimized_resume or ""))
    if step not in ("writer", "ats_judge", "recruiter_judge"):
        raise HttpError(400, "step must be one of: writer, ats_judge, recruiter_judge")
    if not (payload.job_description or "").strip():
        raise HttpError(400, "job_description is required")
    if payload.llm_provider not in LLM_PROVIDERS:
        raise HttpError(400, f"llm_provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")

    # Resolve API key and model
    api_key = None
    config = LLMProviderConfig.objects.filter(provider=payload.llm_provider).first()
    if config and config.encrypted_api_key:
        api_key = decrypt_api_key(config.encrypted_api_key)
    if not api_key:
        raise HttpError(400, f"API key required for {payload.llm_provider}. Connect an API key first.")
    model = payload.llm_model or (config.default_model if config else None) or None
    llm = get_llm(payload.llm_provider, api_key, model)

    from .prompt_store import build_optimizer_graph_prompt_state, get_effective_prompts, save_prompts_to_profile
    from .optimizer_budget import build_optimizer_context_state_raw

    _ov = {}
    if payload.prompt_writer and str(payload.prompt_writer).strip():
        _ov["writer"] = payload.prompt_writer.strip()
    if payload.prompt_ats_judge and str(payload.prompt_ats_judge).strip():
        _ov["ats_judge"] = payload.prompt_ats_judge.strip()
    if payload.prompt_recruiter_judge and str(payload.prompt_recruiter_judge).strip():
        _ov["recruiter_judge"] = payload.prompt_recruiter_judge.strip()
    workflow = None
    if payload.optimizer_workflow_id:
        workflow = OptimizerWorkflow.objects.filter(pk=int(payload.optimizer_workflow_id)).first()
    _graph_prompts = build_optimizer_graph_prompt_state(
        _ov if _ov else None,
        request,
        ats_judge_profile_id=payload.ats_judge_profile_id,
        workflow=workflow,
    )
    # Step-by-step always runs in debug mode
    debug = True
    # Persist edited prompts to UserPromptProfile
    if hasattr(request, "session"):
        merged = get_effective_prompts(request)
        merged.update(
            {
                "writer": payload.prompt_writer or "",
                "recruiter_judge": payload.prompt_recruiter_judge or "",
            }
        )
        save_prompts_to_profile(request, merged)

    def _get_resume_text():
        if file and getattr(file, "size", 0) and getattr(file, "name", "").strip().lower().endswith(PDF_EXTENSION):
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            try:
                return parse_pdf(tmp_path)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
        if payload.use_resume_id:
            ur = get_object_or_404(UserResume, id=payload.use_resume_id)
            return parse_pdf(ur.file.path)
        return None

    jkey = (payload.job_cache_key or "").strip() or None
    try:
        if step == "writer":
            resume_text = _get_resume_text()
            if not resume_text:
                raise HttpError(400, "For writer step provide a resume file upload or use_resume_id.")
            feedback = [f.strip() for f in (payload.feedback or "").split(",") if f.strip()]
            ctx = build_optimizer_context_state_raw(
                resume_text,
                payload.job_description,
                optimization_notes=(payload.optimization_notes or "").strip(),
                pipeline_skills_json=(payload.pipeline_skills_json or "").strip(),
                job_highlights=(payload.job_highlights or "").strip(),
                user_resume_id=int(payload.use_resume_id) if payload.use_resume_id else None,
            )
            state = {
                **ctx,
                "optimized_resume": (payload.optimized_resume or "").strip(),
                "ats_score": 0,
                "recruiter_score": 0,
                "feedback": feedback,
                "iteration_count": 0,
                "llm": llm,
                "job_cache_key": jkey,
                "writer_prompt_template": _graph_prompts["writer_prompt_template"],
                "writer_prompt_system": _graph_prompts["writer_prompt_system"],
                "writer_prompt_user": _graph_prompts["writer_prompt_user"],
                "writer_prompt_legacy": _graph_prompts["writer_prompt_legacy"],
                "ats_judge_prompt_template": _graph_prompts["ats_judge_prompt_template"],
                "ats_judge_prompt_system": _graph_prompts["ats_judge_prompt_system"],
                "ats_judge_prompt_user": _graph_prompts["ats_judge_prompt_user"],
                "ats_judge_prompt_legacy": _graph_prompts["ats_judge_prompt_legacy"],
                "recruiter_judge_prompt_template": _graph_prompts["recruiter_judge_prompt_template"],
                "recruiter_judge_prompt_system": _graph_prompts["recruiter_judge_prompt_system"],
                "recruiter_judge_prompt_user": _graph_prompts["recruiter_judge_prompt_user"],
                "recruiter_judge_prompt_legacy": _graph_prompts["recruiter_judge_prompt_legacy"],
                "debug": debug,
                "max_iterations": 3,
            }
            out = writer_node(state)
            return RunStepResponse(
                step="writer",
                output={
                    "optimized_resume": out.get("optimized_resume", ""),
                    "debug_prompt": out.get("debug_prompt"),
                    "input_tokens": out.get("input_tokens"),
                    "output_tokens": out.get("output_tokens"),
                    "tokens_estimated": out.get("tokens_estimated"),
                    "optimizer_context_budget": ctx.get("optimizer_context_budget"),
                },
            )

        if step in ("ats_judge", "recruiter_judge"):
            optimized_resume = (payload.optimized_resume or "").strip()
            resume_text = ""
            if not optimized_resume:
                parsed = _get_resume_text()
                if parsed:
                    resume_text = (parsed or "").strip()
            if not optimized_resume and not resume_text:
                raise HttpError(
                    400,
                    "For ATS/Recruiter steps provide optimized_resume (current draft text), "
                    "or upload a PDF resume / use_resume_id to score the original resume.",
                )
            state = {
                "resume_text": resume_text,
                "job_description": payload.job_description,
                "optimized_resume": optimized_resume,
                "ats_score": 0,
                "recruiter_score": 0,
                "feedback": [],
                "iteration_count": 0,
                "llm": llm,
                "job_cache_key": jkey,
                "writer_prompt_template": _graph_prompts["writer_prompt_template"],
                "writer_prompt_system": _graph_prompts["writer_prompt_system"],
                "writer_prompt_user": _graph_prompts["writer_prompt_user"],
                "writer_prompt_legacy": _graph_prompts["writer_prompt_legacy"],
                "ats_judge_prompt_template": _graph_prompts["ats_judge_prompt_template"],
                "ats_judge_prompt_system": _graph_prompts["ats_judge_prompt_system"],
                "ats_judge_prompt_user": _graph_prompts["ats_judge_prompt_user"],
                "ats_judge_prompt_legacy": _graph_prompts["ats_judge_prompt_legacy"],
                "recruiter_judge_prompt_template": _graph_prompts["recruiter_judge_prompt_template"],
                "recruiter_judge_prompt_system": _graph_prompts["recruiter_judge_prompt_system"],
                "recruiter_judge_prompt_user": _graph_prompts["recruiter_judge_prompt_user"],
                "recruiter_judge_prompt_legacy": _graph_prompts["recruiter_judge_prompt_legacy"],
                "debug": debug,
                "max_iterations": 3,
            }
            if step == "ats_judge":
                out = ats_judge_node(state)
                fb = out.get("feedback")
                if isinstance(fb, list) and fb:
                    fb = fb[-1]
                elif not isinstance(fb, str):
                    fb = str(fb) if fb else ""
                output = {"ats_score": out.get("ats_score"), "feedback": fb, "debug_prompt": out.get("debug_prompt"), "input_tokens": out.get("input_tokens"), "output_tokens": out.get("output_tokens"), "tokens_estimated": out.get("tokens_estimated")}
                if out.get("last_ats_json") is not None:
                    output["response_json"] = out["last_ats_json"]
                return RunStepResponse(step="ats_judge", output=output)
            else:
                out = recruiter_judge_node(state)
                fb = out.get("feedback")
                if isinstance(fb, list) and fb:
                    fb = fb[-1]
                elif not isinstance(fb, str):
                    fb = str(fb) if fb else ""
                output = {"recruiter_score": out.get("recruiter_score"), "feedback": fb, "debug_prompt": out.get("debug_prompt"), "input_tokens": out.get("input_tokens"), "output_tokens": out.get("output_tokens"), "tokens_estimated": out.get("tokens_estimated")}
                if out.get("last_recruiter_json") is not None:
                    output["response_json"] = out["last_recruiter_json"]
                return RunStepResponse(step="recruiter_judge", output=output)
    except HttpError:
        raise
    except LLMRequestsDisabled as e:
        return RunStepResponse(step=step, output={}, error=str(e))
    except Exception as e:
        err_msg = str(e).strip().replace("\r\n", " ").replace("\r", " ").replace("\n", " ")
        if len(err_msg) > 400:
            err_msg = err_msg[:400] + "..."
        logger.exception("[run_step] step=%s failed: %s", step, e)
        logger.debug("[run_step] traceback: %s", traceback.format_exc())
        return RunStepResponse(step=step, output={}, error=err_msg or "Step failed.")


@router.post("/fit-check", response=FitCheckResponse)
def fit_check(request, payload: FitCheckRequest = Form(...), file: UploadedFile = File(...)):
    """Assess candidate-job fit (0-100). If score < 50, UI should ask user whether to proceed with optimization."""
    _require_api_auth(request)
    if payload.llm_provider not in LLM_PROVIDERS:
        raise HttpError(400, f"llm_provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")
    if len(payload.job_description) > MAX_JOB_DESCRIPTION_LENGTH:
        raise HttpError(400, f"job_description must be at most {MAX_JOB_DESCRIPTION_LENGTH} characters")
    if getattr(file, "size", 0) > MAX_RESUME_FILE_SIZE_BYTES:
        raise HttpError(400, f"Resume file must be at most {MAX_RESUME_FILE_SIZE_BYTES // (1024*1024)} MB")
    if not (getattr(file, "name", "") or "").strip().lower().endswith(PDF_EXTENSION):
        raise HttpError(400, "Resume file must be a PDF")

    api_key = payload.api_key
    if not api_key:
        config = LLMProviderConfig.objects.filter(provider=payload.llm_provider).first()
        if config and config.encrypted_api_key:
            api_key = decrypt_api_key(config.encrypted_api_key)
        if not api_key:
            raise HttpError(400, f"API key required for {payload.llm_provider}.")

    model = payload.llm_model
    if not model:
        config = LLMProviderConfig.objects.filter(provider=payload.llm_provider).first()
        model = config.default_model if config else None
    model = model or None

    llm = get_llm(payload.llm_provider, api_key or None, model)
    prompt_template = payload.prompt_fit_check or None

    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        try:
            resume_text = parse_pdf(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception as e:
        raise HttpError(400, f"Could not read PDF: {e}") from e

    try:
        result = run_fit_check(
            resume_text,
            payload.job_description,
            llm,
            prompt_template,
            usage_query_kind=USAGE_QUERY_API_RESUME_FIT,
        )
    except LLMRequestsDisabled as e:
        raise HttpError(503, str(e)) from e
    return result

@router.get("/prompts", response=PromptsResponse)
def get_prompts(request):
    """Return default LLM prompt templates. Writer also: {full_job_description}, {retrieval_context}, {optimization_notes}, {pipeline_skills_json}, {job_highlights}. JD cleanse: {title}, {job_description}."""
    from .prompts import (
        DEFAULT_ATS_JUDGE_PROMPT,
        DEFAULT_ATS_JUDGE_SYSTEM,
        DEFAULT_ATS_JUDGE_USER,
        DEFAULT_INSIGHTS_PROMPT,
        DEFAULT_INSIGHTS_SYSTEM,
        DEFAULT_INSIGHTS_USER,
        DEFAULT_JD_CLEANSE_PROMPT,
        DEFAULT_JD_CLEANSE_SYSTEM,
        DEFAULT_JD_CLEANSE_USER,
        DEFAULT_MATCHING_PROMPT,
        DEFAULT_MATCHING_SYSTEM,
        DEFAULT_MATCHING_USER,
        DEFAULT_RECRUITER_JUDGE_PROMPT,
        DEFAULT_RECRUITER_JUDGE_SYSTEM,
        DEFAULT_RECRUITER_JUDGE_USER,
        DEFAULT_WRITER_PROMPT,
        DEFAULT_WRITER_SYSTEM,
        DEFAULT_WRITER_USER,
    )
    return {
        "writer": DEFAULT_WRITER_PROMPT,
        "writer_system": DEFAULT_WRITER_SYSTEM,
        "writer_user": DEFAULT_WRITER_USER,
        "ats_judge": DEFAULT_ATS_JUDGE_PROMPT,
        "ats_judge_system": DEFAULT_ATS_JUDGE_SYSTEM,
        "ats_judge_user": DEFAULT_ATS_JUDGE_USER,
        "recruiter_judge": DEFAULT_RECRUITER_JUDGE_PROMPT,
        "recruiter_judge_system": DEFAULT_RECRUITER_JUDGE_SYSTEM,
        "recruiter_judge_user": DEFAULT_RECRUITER_JUDGE_USER,
        "matching": DEFAULT_MATCHING_PROMPT,
        "matching_system": DEFAULT_MATCHING_SYSTEM,
        "matching_user": DEFAULT_MATCHING_USER,
        "insights": DEFAULT_INSIGHTS_PROMPT,
        "insights_system": DEFAULT_INSIGHTS_SYSTEM,
        "insights_user": DEFAULT_INSIGHTS_USER,
        "jd_cleanse": DEFAULT_JD_CLEANSE_PROMPT,
        "jd_cleanse_system": DEFAULT_JD_CLEANSE_SYSTEM,
        "jd_cleanse_user": DEFAULT_JD_CLEANSE_USER,
    }

@router.post("/optimize")
def optimize_resume(request, payload: OptimizeRequest = Form(...), file: UploadedFile = File(...)):
    _require_api_auth(request)
    # Validation
    if payload.llm_provider not in LLM_PROVIDERS:
        raise HttpError(400, f"llm_provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")
    if len(payload.job_description) > MAX_JOB_DESCRIPTION_LENGTH:
        raise HttpError(400, f"job_description must be at most {MAX_JOB_DESCRIPTION_LENGTH} characters")
    if getattr(file, "size", 0) > MAX_RESUME_FILE_SIZE_BYTES:
        raise HttpError(400, f"Resume file must be at most {MAX_RESUME_FILE_SIZE_BYTES // (1024*1024)} MB")
    if not (getattr(file, "name", "") or "").strip().lower().endswith(PDF_EXTENSION):
        raise HttpError(400, "Resume file must be a PDF")
    content_type = getattr(file, "content_type", "") or ""
    if content_type and content_type.lower() not in ALLOWED_CONTENT_TYPES:
        raise HttpError(400, "Resume file must have content type application/pdf")

    # 1. Save Resume (retain original filename for display)
    original_name = (getattr(file, "name", "") or "resume.pdf").strip()
    if original_name and "/" in original_name:
        original_name = original_name.split("/")[-1]
    original_name = (original_name or "resume.pdf")[:255]
    user_resume = UserResume.objects.create(file=file, original_filename=original_name)

    # 2. Save Job Description
    job_desc = JobDescription.objects.create(content=payload.job_description)

    workflow = None
    if payload.optimizer_workflow_id:
        try:
            workflow = OptimizerWorkflow.objects.get(pk=int(payload.optimizer_workflow_id))
        except (OptimizerWorkflow.DoesNotExist, ValueError, TypeError):
            workflow = None

    from .prompt_store import resolve_effective_ats_judge_profile_id, get_ats_judge_profile_by_id

    effective_ats_id = resolve_effective_ats_judge_profile_id(
        ats_judge_profile_id=payload.ats_judge_profile_id,
        workflow=workflow,
    )
    ats_profile = get_ats_judge_profile_by_id(effective_ats_id)

    # 3. Create OptimizedResume record
    optimized = OptimizedResume.objects.create(
        original_resume=user_resume,
        job_description=job_desc,
        status=OptimizedResume.STATUS_QUEUED,
        optimization_notes=(payload.optimization_notes or "").strip(),
        pipeline_skills_json=(payload.pipeline_skills_json or "").strip(),
        job_highlights=(payload.job_highlights or "").strip(),
        optimizer_workflow=workflow,
        ats_judge_profile=ats_profile,
    )

    # Resolve API key: request > stored config > env
    api_key = payload.api_key
    if not api_key:
        config = LLMProviderConfig.objects.filter(provider=payload.llm_provider).first()
        if config and config.encrypted_api_key:
            api_key = decrypt_api_key(config.encrypted_api_key)
        if not api_key:
            raise HttpError(400, f"API key required for {payload.llm_provider}. Enter key in LLM Config and connect, or pass api_key.")

    model = payload.llm_model
    if not model:
        config = LLMProviderConfig.objects.filter(provider=payload.llm_provider).first()
        model = config.default_model if config else None
    model = model or None

    prompts = None
    if payload.prompt_writer or payload.prompt_ats_judge or payload.prompt_recruiter_judge:
        prompts = {
            "writer": payload.prompt_writer or None,
            "ats_judge": payload.prompt_ats_judge or None,
            "recruiter_judge": payload.prompt_recruiter_judge or None,
        }

    # Debug flag: form sends "true"/"false" as string; read from POST so it's not lost
    debug_val = request.POST.get("debug", "") if hasattr(request, "POST") else ""
    debug = str(debug_val).lower() in ("true", "1", "on")

    rate_limit_delay = 0
    try:
        rld = request.POST.get("rate_limit_delay", "") if hasattr(request, "POST") else ""
        if rld not in ("", None):
            rate_limit_delay = float(rld)
    except (TypeError, ValueError):
        pass

    max_iterations = 3
    try:
        mi = request.POST.get("max_iterations", "") if hasattr(request, "POST") else ""
        if mi not in ("", None):
            max_iterations = int(mi)
    except (TypeError, ValueError):
        pass

    # Parse optional workflow step order (JSON array string)
    workflow_steps = None
    if payload.workflow_steps and (payload.workflow_steps or "").strip():
        try:
            raw = json.loads(payload.workflow_steps)
            if isinstance(raw, list) and raw:
                invalid = [s for s in raw if not isinstance(s, str) or s not in VALID_STEP_IDS]
                if invalid:
                    raise HttpError(400, f"Invalid workflow_steps: {invalid}. Allowed: {sorted(VALID_STEP_IDS)}")
                workflow_steps = raw
        except json.JSONDecodeError as e:
            raise HttpError(400, f"workflow_steps must be a JSON array of step ids: {e}") from e
    loop_to = (payload.loop_to or "").strip() or None
    if loop_to is not None and loop_to not in VALID_STEP_IDS:
        raise HttpError(400, f"loop_to must be one of: {sorted(VALID_STEP_IDS)}")

    score_threshold = 85
    if payload.score_threshold is not None:
        st = int(payload.score_threshold)
        if not 0 <= st <= 100:
            raise HttpError(400, "score_threshold must be between 0 and 100")
        score_threshold = st

    # Enqueue Huey task (Redis-backed worker)
    loop_to_val = loop_to
    max_it_val = max_iterations
    score_t_val = score_threshold
    if workflow:
        if not workflow_steps and workflow.steps:
            workflow_steps = [
                s for s in workflow.steps if isinstance(s, str) and s in VALID_STEP_IDS
            ] or None
        if not loop_to_val and (workflow.loop_to or "").strip():
            loop_to_val = workflow.loop_to.strip()
        if workflow.max_iterations:
            max_it_val = max(1, min(int(workflow.max_iterations), 5))
        if workflow.score_threshold is not None:
            score_t_val = max(0, min(100, int(workflow.score_threshold)))

    result = optimize_resume_task(
        optimized.id,
        job_desc.id,
        payload.llm_provider,
        api_key or "",
        model,
        prompts=prompts,
        debug=debug,
        rate_limit_delay=rate_limit_delay,
        max_iterations=max_it_val,
        score_threshold=score_t_val,
        workflow_steps=workflow_steps,
        loop_to=loop_to_val,
        ats_judge_profile_id=effective_ats_id,
    )
    task_id = result.id if result else None
    return {"task_id": task_id, "resume_id": optimized.id}

def get_status_data(resume_id: int):
    optimized = get_object_or_404(OptimizedResume, id=resume_id)
    logs = AgentLog.objects.filter(optimized_resume=optimized).order_by('created_at')

    # Use stored scores; fall back to last log for backward compat
    ats_score = optimized.ats_score
    recruiter_score = optimized.recruiter_score
    if ats_score is None or recruiter_score is None:
        for log in logs:
            if 'ats_score' in log.thought:
                ats_score = log.thought['ats_score']
            if 'recruiter_score' in log.thought:
                recruiter_score = log.thought['recruiter_score']

    # "status" must stay the canonical workflow value (queued|running|completed|failed) so
    # the optimizer UI and templates can branch correctly. Human progress lives in status_display.
    return {
        "status": optimized.status,
        "status_display": (optimized.status_display or "").strip() or None,
        "ats_score": ats_score,
        "recruiter_score": recruiter_score,
        "optimized_content": optimized.optimized_content,
        "error_message": optimized.error_message,
        "total_input_tokens": optimized.total_input_tokens,
        "total_output_tokens": optimized.total_output_tokens,
        "logs": [{"step": l.step_name, "thought": l.thought} for l in logs],
        "optimizer_context_snapshot": optimized.optimizer_context_snapshot,
    }

@router.get("/status/{resume_id}", response=StatusResponse)
def get_status(request, resume_id: int):
    _require_api_auth(request)
    return get_status_data(resume_id)


CANCELLED_MESSAGE = "Cancelled by user"


class SaveDraftRequest(Schema):
    optimized_content: str


class SaveDraftResponse(Schema):
    ok: bool = True
    optimized_content: str


@router.post("/status/{resume_id}/draft", response=SaveDraftResponse)
def save_optimized_draft(request, resume_id: int, payload: SaveDraftRequest):
    """Save manual edits to the optimized resume draft before PDF/Word export."""
    from .services import DraftSaveError, save_optimized_draft_content

    _require_api_auth(request)
    try:
        optimized = save_optimized_draft_content(resume_id, payload.optimized_content or "")
    except DraftSaveError as e:
        raise HttpError(e.status_code, e.message) from e
    return {"ok": True, "optimized_content": optimized.optimized_content or ""}


@router.post("/status/{resume_id}/cancel")
def cancel_optimization(request, resume_id: int):
    """Mark a running or queued optimization as cancelled. The background task will stop on its next check."""
    _require_api_auth(request)
    optimized = get_object_or_404(OptimizedResume, id=resume_id)
    if optimized.status not in (OptimizedResume.STATUS_QUEUED, OptimizedResume.STATUS_RUNNING):
        return {"success": False, "message": "Optimization is not running or queued."}
    optimized.status = OptimizedResume.STATUS_FAILED
    optimized.error_message = CANCELLED_MESSAGE
    optimized.status_display = ""
    optimized.save(update_fields=["status", "error_message", "status_display"])
    return {"success": True, "message": "Optimization cancelled."}


# --- Workflows (saved custom optimizer workflows) ---

class WorkflowSchema(Schema):
    id: int
    name: str
    steps: List[str]
    loop_to: str
    max_iterations: int
    score_threshold: int
    ats_judge_profile_id: Optional[int] = None


class WorkflowCreateUpdate(Schema):
    name: str
    steps: List[str]
    loop_to: Optional[str] = ""
    max_iterations: Optional[int] = 3
    score_threshold: Optional[int] = 85
    ats_judge_profile_id: Optional[int] = None


class AtsJudgeProfileListItem(Schema):
    id: int
    name: str
    slug: str
    is_builtin: bool
    is_default: bool


class AtsJudgeProfileDetail(Schema):
    id: int
    name: str
    slug: str
    is_builtin: bool
    is_default: bool
    ats_judge: str
    ats_judge_system: str
    ats_judge_user: str
    ats_judge_legacy_combined: str = ""


class AtsJudgeProfileCreateUpdate(Schema):
    name: str
    slug: Optional[str] = None
    is_default: Optional[bool] = False
    ats_judge: Optional[str] = ""
    ats_judge_system: Optional[str] = ""
    ats_judge_user: Optional[str] = ""
    ats_judge_legacy_combined: Optional[str] = ""


def _workflow_to_schema(w: OptimizerWorkflow) -> dict:
    return {
        "id": w.id,
        "name": w.name,
        "steps": w.steps,
        "loop_to": w.loop_to or "",
        "max_iterations": w.max_iterations,
        "score_threshold": w.score_threshold,
        "ats_judge_profile_id": w.ats_judge_profile_id,
    }


def _ats_profile_to_detail(p: AtsJudgeProfile) -> dict:
    from .prompt_store import get_ats_judge_profile_display

    disp = get_ats_judge_profile_display(p)
    return {
        "id": p.id,
        "name": p.name,
        "slug": p.slug,
        "is_builtin": p.is_builtin,
        "is_default": p.is_default,
        "ats_judge": disp["ats_judge"],
        "ats_judge_system": disp["ats_judge_system"],
        "ats_judge_user": disp["ats_judge_user"],
        "ats_judge_legacy_combined": disp["ats_judge_legacy_combined"],
    }


def _resolve_ats_profile_fk(profile_id: Optional[int]) -> Optional[AtsJudgeProfile]:
    if profile_id is None:
        return None
    try:
        return AtsJudgeProfile.objects.get(pk=int(profile_id))
    except (AtsJudgeProfile.DoesNotExist, ValueError, TypeError):
        raise HttpError(400, f"Invalid ats_judge_profile_id: {profile_id}")


def _validate_workflow_steps(steps: list) -> None:
    if not steps:
        raise HttpError(400, "steps must not be empty")
    invalid = [s for s in steps if not isinstance(s, str) or s not in VALID_STEP_IDS]
    if invalid:
        raise HttpError(400, f"Invalid step id(s): {invalid}. Allowed: {sorted(VALID_STEP_IDS)}")


@router.get("/ats-judge-profiles", response=List[AtsJudgeProfileListItem])
def list_ats_judge_profiles_api(request):
    """List named ATS judge prompt profiles."""
    _require_api_auth(request)
    return [
        {
            "id": p.id,
            "name": p.name,
            "slug": p.slug,
            "is_builtin": p.is_builtin,
            "is_default": p.is_default,
        }
        for p in AtsJudgeProfile.objects.all().order_by("name")
    ]


@router.get("/ats-judge-profiles/{profile_id}", response=AtsJudgeProfileDetail)
def get_ats_judge_profile_api(request, profile_id: int):
    _require_api_auth(request)
    p = get_object_or_404(AtsJudgeProfile, id=profile_id)
    return _ats_profile_to_detail(p)


@router.post("/ats-judge-profiles", response=AtsJudgeProfileDetail)
def create_ats_judge_profile_api(request, payload: AtsJudgeProfileCreateUpdate):
    _require_api_auth(request)
    from .prompt_store import save_ats_judge_profile

    name = (payload.name or "").strip()
    if not name:
        raise HttpError(400, "name is required")
    slug = (payload.slug or "").strip() or None
    if slug and AtsJudgeProfile.objects.filter(slug=slug).exists():
        raise HttpError(400, f"slug already exists: {slug}")
    leg = (payload.ats_judge_legacy_combined or payload.ats_judge or "").strip()
    p = AtsJudgeProfile(name=name, slug=slug or "", is_default=bool(payload.is_default))
    if slug:
        p.slug = slug
    save_ats_judge_profile(
        p,
        ats_judge=leg,
        ats_judge_system=(payload.ats_judge_system or "").strip(),
        ats_judge_user=(payload.ats_judge_user or "").strip(),
    )
    if payload.is_default:
        p.is_default = True
        p.save()
    return _ats_profile_to_detail(p)


@router.put("/ats-judge-profiles/{profile_id}", response=AtsJudgeProfileDetail)
def update_ats_judge_profile_api(request, profile_id: int, payload: AtsJudgeProfileCreateUpdate):
    _require_api_auth(request)
    from .prompt_store import save_ats_judge_profile

    p = get_object_or_404(AtsJudgeProfile, id=profile_id)
    name = (payload.name or "").strip()
    if not name:
        raise HttpError(400, "name is required")
    slug = (payload.slug or "").strip()
    if slug and slug != p.slug and AtsJudgeProfile.objects.filter(slug=slug).exclude(pk=p.pk).exists():
        raise HttpError(400, f"slug already exists: {slug}")
    if slug:
        p.slug = slug
    leg = (payload.ats_judge_legacy_combined or payload.ats_judge or "").strip()
    save_ats_judge_profile(
        p,
        name=name,
        ats_judge=leg,
        ats_judge_system=(payload.ats_judge_system or "").strip(),
        ats_judge_user=(payload.ats_judge_user or "").strip(),
    )
    if payload.is_default is not None:
        p.is_default = bool(payload.is_default)
        p.save()
    return _ats_profile_to_detail(p)


@router.delete("/ats-judge-profiles/{profile_id}")
def delete_ats_judge_profile_api(request, profile_id: int):
    _require_api_auth(request)
    p = get_object_or_404(AtsJudgeProfile, id=profile_id)
    if p.is_builtin and AtsJudgeProfile.objects.count() <= 1:
        raise HttpError(400, "Cannot delete the only built-in ATS profile.")
    was_default = p.is_default
    p.delete()
    if was_default:
        fallback = AtsJudgeProfile.objects.order_by("pk").first()
        if fallback:
            fallback.is_default = True
            fallback.save()
    return {"success": True}


@router.get("/workflows", response=List[WorkflowSchema])
def list_workflows(request):
    """List all saved optimizer workflows."""
    return [_workflow_to_schema(w) for w in OptimizerWorkflow.objects.all()]


@router.get("/workflows/{workflow_id}", response=WorkflowSchema)
def get_workflow(request, workflow_id: int):
    w = get_object_or_404(OptimizerWorkflow, id=workflow_id)
    return _workflow_to_schema(w)


@router.post("/workflows", response=WorkflowSchema)
def create_workflow_api(request, payload: WorkflowCreateUpdate):
    _validate_workflow_steps(payload.steps)
    loop_to = (payload.loop_to or "").strip() or ""
    if loop_to and loop_to not in VALID_STEP_IDS:
        raise HttpError(400, f"loop_to must be one of: {sorted(VALID_STEP_IDS)} or empty")
    max_it = payload.max_iterations if payload.max_iterations is not None else 3
    max_it = max(1, min(5, max_it))
    score_t = payload.score_threshold if payload.score_threshold is not None else 85
    score_t = max(0, min(100, score_t))
    w = OptimizerWorkflow.objects.create(
        name=payload.name.strip(),
        steps=payload.steps,
        loop_to=loop_to,
        max_iterations=max_it,
        score_threshold=score_t,
        ats_judge_profile=_resolve_ats_profile_fk(payload.ats_judge_profile_id),
    )
    return _workflow_to_schema(w)


@router.put("/workflows/{workflow_id}", response=WorkflowSchema)
def update_workflow(request, workflow_id: int, payload: WorkflowCreateUpdate):
    w = get_object_or_404(OptimizerWorkflow, id=workflow_id)
    _validate_workflow_steps(payload.steps)
    loop_to = (payload.loop_to or "").strip() or ""
    if loop_to and loop_to not in VALID_STEP_IDS:
        raise HttpError(400, f"loop_to must be one of: {sorted(VALID_STEP_IDS)} or empty")
    max_it = payload.max_iterations if payload.max_iterations is not None else 3
    max_it = max(1, min(5, max_it))
    score_t = payload.score_threshold if payload.score_threshold is not None else 85
    score_t = max(0, min(100, score_t))
    w.name = payload.name.strip()
    w.steps = payload.steps
    w.loop_to = loop_to
    w.max_iterations = max_it
    w.score_threshold = score_t
    w.ats_judge_profile = (
        _resolve_ats_profile_fk(payload.ats_judge_profile_id)
        if payload.ats_judge_profile_id
        else None
    )
    w.save()
    return _workflow_to_schema(w)


@router.delete("/workflows/{workflow_id}")
def delete_workflow(request, workflow_id: int):
    w = get_object_or_404(OptimizerWorkflow, id=workflow_id)
    w.delete()
    return {"success": True}


# --- LLM config: connect (validate + save key), list models ---

class ConnectRequest(Schema):
    provider: str
    api_key: str


@router.post("/llm/connect")
def llm_connect(request, payload: ConnectRequest):
    """Validate API key by listing models; on success save key (encrypted) and return models."""
    if payload.provider not in LLM_PROVIDERS:
        raise HttpError(400, f"provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")
    if not payload.api_key or not payload.api_key.strip():
        raise HttpError(400, "api_key is required")
    try:
        models = list_models_for_provider(payload.provider, payload.api_key.strip())
    except Exception as e:
        err = str(e)
        if is_auth_error(e):
            LLMProviderConfig.objects.filter(provider=payload.provider).delete()
            raise HttpError(401, f"Invalid API key: {err}")
        raise HttpError(400, err)
    config, _ = LLMProviderConfig.objects.get_or_create(provider=payload.provider, defaults={"encrypted_api_key": ""})
    config.encrypted_api_key = encrypt_api_key(payload.api_key.strip())
    config.last_validated_at = timezone.now()
    config.save(update_fields=["encrypted_api_key", "last_validated_at", "updated_at"])
    return {"success": True, "models": models}


@router.get("/llm/models")
def llm_models(request, provider: str):
    """Return list of models for provider using stored API key or env fallback. 401 if no key or key invalid."""
    if provider not in LLM_PROVIDERS:
        raise HttpError(400, f"provider must be one of: {', '.join(sorted(LLM_PROVIDERS))}")
    config = LLMProviderConfig.objects.filter(provider=provider).first()
    api_key = None
    if config and config.encrypted_api_key:
        api_key = decrypt_api_key(config.encrypted_api_key)
        if not api_key:
            config.encrypted_api_key = ""
            config.last_validated_at = None
            config.save(update_fields=["encrypted_api_key", "last_validated_at", "updated_at"])
            raise HttpError(401, "Stored key invalid. Please enter a new key and Connect.")
    if not api_key:
        from .pipeline_llm_skill_extract import resolve_provider_api_key

        api_key = resolve_provider_api_key(provider)
    if not api_key:
        raise HttpError(401, "No API key stored. Enter key and click Connect, or set env for this provider.")
    had_stored_key = bool(config and config.encrypted_api_key)
    try:
        models = list_models_for_provider(provider, api_key)
        if config:
            config.last_validated_at = timezone.now()
            config.save(update_fields=["last_validated_at", "updated_at"])
        default_m = (config.default_model if config else None) or DEFAULT_MODELS.get(provider)
        return {"models": models, "default_model": default_m}
    except Exception as e:
        if is_auth_error(e):
            if had_stored_key and config:
                config.encrypted_api_key = ""
                config.last_validated_at = None
                config.save(update_fields=["encrypted_api_key", "last_validated_at", "updated_at"])
            raise HttpError(401, f"API key no longer valid: {e}")
        raise HttpError(400, str(e))


@router.post("/llm/set-default-model")
def llm_set_default_model(request, provider: str, model: str):
    """Set default model for provider (stored config)."""
    if provider not in LLM_PROVIDERS:
        raise HttpError(400, "Invalid provider")
    config = LLMProviderConfig.objects.filter(provider=provider).first()
    if not config:
        raise HttpError(404, "Connect with an API key first")
    config.default_model = model
    config.save(update_fields=["default_model", "updated_at"])
    return {"success": True}


def _parse_markdown_blocks(content: str):
    """Yield (block_type, text). block_type: heading1, heading2, bullet, paragraph."""
    if not content:
        return
    for line in (content or "").replace("\r", "").split("\n"):
        raw = line
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            yield ("heading1", line[2:].strip())
        elif line.startswith("## "):
            yield ("heading2", line[3:].strip())
        elif line.startswith("- ") or line.startswith("* "):
            yield ("bullet", line[2:].strip())
        else:
            yield ("paragraph", raw.strip())


def _normalize_export_content(text: str) -> str:
    """Normalize common Unicode punctuation and whitespace for PDF/Word export."""
    if not text:
        return text
    import unicodedata

    text = unicodedata.normalize("NFKC", text)
    replacements = {
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2015": "-",
        "\u2212": "-",
        "\u00A0": " ",
        "\u202F": " ",
        "\u2007": " ",
        "\u2009": " ",
        "\u200B": "",
        "\u200C": "",
        "\u200D": "",
        "\u2018": "'",
        "\u2019": "'",
        "\u201C": '"',
        "\u201D": '"',
        "\u2026": "...",
    }
    return text.translate(str.maketrans(replacements))


def _split_style_spans(text: str):
    """Yield (is_bold, is_italic, segment) for text with **bold** and *italic* markers."""
    import re
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("***") and p.endswith("***"):
            yield (True, True, p[3:-3])
        elif p.startswith("**") and p.endswith("**"):
            yield (True, False, p[2:-2])
        elif p.startswith("*") and p.endswith("*"):
            yield (False, True, p[1:-1])
        else:
            yield (False, False, p)


def _draw_rich_text(
    c,
    x: float,
    y: float,
    max_width: float,
    segments,
    line_height: float,
    font_size: float = 10,
    first_line_indent: float = None,
    default_font_name: str = "Times-Roman",
    text_color = None,
) -> float:
    if line_height is None:
        line_height = font_size * 1.5
    if first_line_indent is None:
        first_line_indent = x
    current_x = x
    current_y = y
    line_start = x
    for is_bold, is_italic, segment in segments:
        import re
        tokens = re.findall(r"\s+|\S+", segment)
        for token in tokens:
            if is_bold and is_italic:
                # Map default font to correct variant
                if "Times" in default_font_name:
                    font_name = "Times-BoldItalic"
                else:
                    font_name = "Helvetica-BoldOblique"
            elif is_bold:
                if "Times" in default_font_name:
                    font_name = "Times-Bold"
                else:
                    font_name = "Helvetica-Bold"
            elif is_italic:
                if "Times" in default_font_name:
                    font_name = "Times-Italic"
                else:
                    font_name = "Helvetica-Oblique"
            else:
                font_name = default_font_name
            c.setFont(font_name, font_size)
            if text_color:
                c.setFillColor(text_color)
            token_width = c.stringWidth(token, font_name, font_size)
            if token.strip() and current_x + token_width > line_start + max_width:
                current_y -= line_height
                current_x = first_line_indent
                line_start = first_line_indent
            if token.isspace() and current_x == line_start:
                continue
            c.drawString(current_x, current_y, token)
            current_x += token_width
    # Reset color to black for next section
    from reportlab.lib import colors
    c.setFillColor(colors.black)
    return current_y - line_height


def _build_export_pdf(content: str) -> io.BytesIO:
    """Build PDF from markdown-style content with headings, bullets, bold, italic.
    Uses Times-Roman for professional typography and blue headings for visual hierarchy."""
    content = _normalize_export_content(content)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        y = height - inch
        body_line_height = 18
        text_width = width - 2 * inch
        heading_color = (31/255, 78/255, 121/255)
        
        for block_type, text in _parse_markdown_blocks(content):
            block_height = body_line_height
            if block_type == "heading1":
                block_height = 26
            elif block_type == "heading2":
                block_height = 22
            elif block_type == "bullet":
                block_height = 20
            
            if y < inch + block_height * 3:
                c.showPage()
                y = height - inch
            
            if block_type == "heading1":
                y -= body_line_height / 4
                y = _draw_rich_text(
                    c,
                    inch,
                    y,
                    text_width,
                    _split_style_spans(text),
                    block_height,
                    font_size=18,
                    default_font_name="Times-Bold",
                    text_color=heading_color,
                )
                y -= body_line_height / 3
            elif block_type == "heading2":
                y -= body_line_height / 4
                y = _draw_rich_text(
                    c,
                    inch,
                    y,
                    text_width,
                    _split_style_spans(text),
                    block_height,
                    font_size=13,
                    default_font_name="Times-Bold",
                    text_color=heading_color,
                )
                y -= body_line_height / 3
            elif block_type == "bullet":
                c.setFont("Times-Roman", 11)
                c.setFillColor(colors.black)
                c.drawString(inch, y, "•")
                y = _draw_rich_text(
                    c,
                    inch + 14,
                    y,
                    text_width - 14,
                    _split_style_spans(text),
                    block_height,
                    font_size=11,
                    first_line_indent=inch + 14,
                    default_font_name="Times-Roman",
                )
                y -= body_line_height / 6
            else:
                y = _draw_rich_text(
                    c,
                    inch,
                    y,
                    text_width,
                    _split_style_spans(text),
                    body_line_height,
                    font_size=11,
                    default_font_name="Times-Roman",
                )
                y -= body_line_height / 5
        
        c.save()
        buf.seek(0)
        return buf
    except ImportError:
        return None


def _build_export_docx(content: str) -> io.BytesIO:
    """Build Word from markdown-style content with headings, bullets, bold, italic."""
    content = _normalize_export_content(content)
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for block_type, text in _parse_markdown_blocks(content):
            if block_type == "heading1":
                p = doc.add_heading(text, level=0)
                p.paragraph_format.space_after = Pt(8)
            elif block_type == "heading2":
                p = doc.add_heading(text, level=1)
                p.paragraph_format.space_after = Pt(6)
            elif block_type == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                for is_bold, is_italic, segment in _split_style_spans(text):
                    r = p.add_run(segment)
                    r.bold = is_bold
                    r.italic = is_italic
                p.paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph()
                for is_bold, is_italic, segment in _split_style_spans(text):
                    r = p.add_run(segment)
                    r.bold = is_bold
                    r.italic = is_italic
                p.paragraph_format.space_after = Pt(6)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except ImportError:
        return None


def _apply_export_replacements(content: str, request) -> str:
    if not content:
        return content
    replacements = []
    params = request.GET
    for key, value in params.items():
        if not key.startswith("replace_token_") or not value:
            continue
        suffix = key[len("replace_token_"):]
        replacement = params.get(f"replace_value_{suffix}")
        if replacement is None or replacement == "":
            continue
        replacements.append((value, replacement))

    if not replacements:
        raw_replacements = request.session.get("export_replacements") or []
        for entry in raw_replacements:
            if not isinstance(entry, dict):
                continue
            token = (entry.get("token") or "").strip()
            replacement = entry.get("value")
            if not token or replacement is None or replacement == "":
                continue
            replacements.append((token, replacement))

    for token, replacement in replacements:
        content = content.replace(token, replacement)
    return content


@router.get("/export/{resume_id}/pdf")
def export_pdf(request, resume_id: int):
    """Export optimized resume as PDF. Returns 404 if not found or not completed."""
    optimized = get_object_or_404(OptimizedResume, id=resume_id)
    if optimized.status != OptimizedResume.STATUS_COMPLETED or not optimized.optimized_content:
        raise HttpError(404, "Optimized resume not ready for export")
    content = _apply_export_replacements(optimized.optimized_content, request)
    buf = _build_export_pdf(content)
    if buf is None:
        raise HttpError(503, "PDF export requires reportlab; install with: pip install reportlab")
    return FileResponse(buf, as_attachment=True, filename="optimized_resume.pdf", content_type="application/pdf")


@router.get("/export/{resume_id}/docx")
def export_docx(request, resume_id: int):
    """Export optimized resume as Word. Returns 404 if not found or not completed."""
    optimized = get_object_or_404(OptimizedResume, id=resume_id)
    if optimized.status != OptimizedResume.STATUS_COMPLETED or not optimized.optimized_content:
        raise HttpError(404, "Optimized resume not ready for export")
    content = _apply_export_replacements(optimized.optimized_content, request)
    buf = _build_export_docx(content)
    if buf is None:
        raise HttpError(503, "Word export requires python-docx; install with: pip install python-docx")
    return FileResponse(buf, as_attachment=True, filename="optimized_resume.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
